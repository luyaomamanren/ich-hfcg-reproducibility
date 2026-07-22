#!/usr/bin/env python3
"""Revise Frontiers cover letter, scope statement and figure alt text."""
from pathlib import Path
from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "outputs" / "Frontiers_submission_package"


def revise(source_name, target_name, replacements):
    doc = Document(OUT / source_name)
    for idx, text in replacements.items():
        doc.paragraphs[idx].text = text
    doc.save(OUT / target_name)
    print(OUT / target_name)


title = ("Raw-data transcriptomic reanalysis and genetic sensitivity analyses define an "
         "expression-associated myeloid injury program in human intracerebral hemorrhage")

revise("Cover_Letter_Frontiers_Inflammation.docx", "Cover_Letter_Frontiers_Inflammation_revised.docx", {
    0: "22 July 2026",
    3: f"Please consider our manuscript entitled \u201c{title}\u201d for publication as an Original Research article in the Inflammation section of Frontiers in Immunology.",
    4: ("Secondary injury after intracerebral hemorrhage involves inflammatory, hypoxic, and iron-handling responses, but expression changes alone cannot identify causal therapeutic targets. We therefore reconstructed the principal GEO cohorts from their original CEL, raw-count, and 10x deposits; fitted sample-aware bulk and single-cell models; and re-ran the genetic analysis using formal 1000 Genomes European LD clumping, instrument-strength statistics, Steiger directionality, multiple-testing correction, and regional colocalization."),
    5: ("The revised analysis provides three advances. First, it supplies a transparent, fully auditable raw-data workflow with complete code, intermediate results, environment records, and checksums. Second, it reproduces an expression- and cell-state-associated ICH injury program while distinguishing biological association from causation. Third, it corrects the genetic interpretation: no evaluated gene survived the corrected four-gene MR analysis, and available colocalization results did not support a shared signal. TLR4 and STAT3 are therefore described only as genetically prioritized candidate regulators requiring validation; HMOX1 and NLRP3 are downstream expression- and cell-state-associated candidates."),
    6: ("The manuscript fits the Inflammation section because it examines human innate immune activation and inflammatory cell states after ICH while directly addressing the evidentiary standards needed for regulator prioritization. Its main contribution is not a therapeutic claim, but a reproducible integration of bulk, single-cell, and genetic evidence with explicit limits on causal inference."),
})

revise("Scope_Statement.docx", "Scope_Statement_revised.docx", {
    2: ("This Original Research manuscript is directly relevant to the Inflammation section because it investigates innate immune activation and inflammatory cell states during human intracerebral hemorrhage. The principal GEO cohorts were reconstructed from original CEL, raw-count, and 10x deposits, and the genetic component was re-evaluated with LD clumping, instrument-strength assessment, Steiger directionality, multiple-testing correction, and regional colocalization. The study supports an expression-associated myeloid injury program while showing why expression and cell-state prominence should not be equated with causal target validation. TLR4 and STAT3 are retained only as genetically prioritized candidate regulators requiring independent validation; HMOX1 and NLRP3 are interpreted as downstream expression- and cell-state-associated candidates. This combination of human inflammation biology, transparent statistical sensitivity analyses, and reproducible data processing is closely aligned with the section's focus on inflammatory mechanisms and signaling in disease progression."),
    3: "Word count: 124",
})

revise("Figure_Alt_Text.docx", "Figure_Alt_Text_revised.docx", {
    3: "Flow diagram separating raw-data reconstruction, transcriptome-wide expression analysis, sample-aware single-cell localization, and corrected genetic causal-prioritization analyses.",
    5: "Discovery figure summarizing raw-CEL quality control, patient-blocked differential expression, and descriptive inflammatory pathway context in GSE24265.",
    7: "Validation figure summarizing raw-count processing, technical-replicate collapse, and candidate estimates from the complete patient-blocked GSE163256 model.",
    9: "Two-panel corrected genetic figure. Panel A shows Wald-ratio log-odds effect estimates and 95% confidence intervals after LD clumping, centered on a zero null line. Panel B shows colocalization PP.H4 values for NLRP3 and TLR4 and labels STAT3 and HMOX1 as not estimable in the selected exposure dataset.",
    13: "Single-cell figure showing major immune cell classes identified from the nine raw GSE266873 10x matrices and candidate-expression localization across post-ICH time groups.",
    15: "Exploratory myeloid-state views retained for hypothesis generation; these panels are not interpreted as causal evidence.",
    17: "Exploratory cell-cell communication views retained as contextual material; they do not support target-validation claims.",
    25: "Complete MR audit showing harmonized and clumped instruments, F statistics, Steiger directionality, multiple-testing-adjusted estimates, and explicit non-estimability of multi-instrument diagnostics after one variant remained per gene.",
})

supp = Document(OUT / "Supplementary_Material_Figures_and_Table_Legends.docx")
supp.paragraphs[1].text = title
supp.paragraphs[14].clear()
supp.paragraphs[14].add_run().add_picture(
    str(ROOT / "results" / "figures" / "Supplementary_Figure_S4_corrected_preview.png"),
    width=Inches(6.5),
)
supp.paragraphs[15].text = ("Supplementary Figure S4. Corrected MR instrument audit. "
    "(A) Minimum F statistic for the single independent variant retained per gene after "
    "1000 Genomes European LD clumping; the dashed line denotes F=10. (B) Estimability "
    "summary showing that Cochran Q and MR-Egger intercept tests are not estimable with "
    "one instrument, whereas Steiger directionality was correct for all four genes.")
supp.paragraphs[31].text = ("Supplementary Table S5. Complete MR audit containing candidate "
    "instruments before and after LD clumping, harmonized alleles and effects, per-variant "
    "F statistics, Wald-ratio estimates, Benjamini-Hochberg-adjusted P values, Steiger "
    "directionality, and explicit estimability notes for heterogeneity and pleiotropy tests.")
supp.paragraphs[32].text = ("Supplementary Table S6. Regional colocalization results and "
    "p12 prior-sensitivity analysis, including explicit non-estimability where the selected "
    "whole-blood exposure dataset contained no regional variants.")
supp_target = OUT / "Supplementary_Material_Figures_and_Table_Legends_revised.docx"
supp.save(supp_target)
print(supp_target)
